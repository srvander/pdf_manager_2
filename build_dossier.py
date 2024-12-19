import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx2pdf import convert
from PyPDF2 import PdfWriter, PdfReader
import PyPDF2
from PIL import Image, ImageTk
import os
from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.lib.pagesizes import letter
import os
from docx import Document  # Import Document for Word file creation
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Import WD_ALIGN_PARAGRAPH for paragraph alignment
from custom_widgets import CustomFrame, CustomLabel, CustomButton_A, CustomButton_B
import customtkinter as ctk
from CTkListbox import CTkListbox


class BuildDossier(CustomFrame):
    # ...existing code...

    def __init__(self, master):
        super().__init__(master)
        # ...existing code...
        
        # Create top frame for buttons
        top_frame = CustomFrame(self)
        top_frame.pack(pady=10)
        
        # Add "Seleccionar Archivos" and "Generar Expediente" buttons to top_frame
        select_button = CustomButton_A(top_frame, text="SELECCIONAR ARCHIVOS", command=self.select_files)
        select_button.pack(side=tk.LEFT, padx=5)
        
        generar_button = CustomButton_A(top_frame, text="GENERAR EXPEDIENTE", command=lambda: self.generar_expediente(self.selected_files))
        generar_button.pack(side=tk.LEFT, padx=5)
        
        self.selected_files_label = CustomLabel(self, text="No se han seleccionado archivos.")
        self.selected_files_label.pack(pady=5)

        # Add checkbox for foliar option
        self.foliar_var = tk.BooleanVar()
        self.start_page_var = tk.IntVar(value=1)  # Variable for initial page number

        # Create a frame for foliar options
        foliar_frame = CustomFrame(self)
        foliar_frame.pack(pady=5)

        foliar_checkbox = ctk.CTkCheckBox(foliar_frame, text="Foliar Expediente", variable=self.foliar_var, command=self.toggle_start_page_entry)
        foliar_checkbox.pack(side=tk.LEFT)

        start_page_label = CustomLabel(foliar_frame, text="Número inicial:")
        start_page_label.pack(side=tk.LEFT, padx=5)

        self.start_page_entry = ttk.Entry(foliar_frame, textvariable=self.start_page_var, width=5, state='disabled')
        self.start_page_entry.pack(side=tk.LEFT)

        frame = CustomFrame(self)
        frame.pack(pady=10, fill=tk.BOTH, expand=True)
        
        self.file_listbox = CTkListbox(frame,
                                       font=("Montserrat bold", 10),
                                       text_color="black",)
        self.file_listbox.pack(side="left", fill=tk.BOTH, expand=True, padx=10)
        
        button_frame = CustomFrame(frame)
        button_frame.pack(side=tk.LEFT, padx=5)
        
        self.move_up_button = CustomButton_B(button_frame, text="▲", command=self.move_up, state=tk.DISABLED)
        self.move_up_button.pack(pady=5)
        
        self.move_down_button = CustomButton_B(button_frame, text="▼", command=self.move_down, state=tk.DISABLED)
        self.move_down_button.pack(pady=5)
        
        # Create bottom frame
        bottom_frame = CustomFrame(self)
        bottom_frame.pack(pady=10)
        
        self.back_button = CustomButton_A(bottom_frame, text="VOLVER AL MENÚ", command=self.volver_menu_principal)
        self.back_button.pack()

    def select_files(self):
        file_types = [
            ("Todos los archivos soportados", "*.pdf *.docx *.jpg *.jpeg *.png"),
            ("Archivos PDF", "*.pdf"),
            ("Archivos Word", "*.docx"),
            ("Imágenes", "*.jpg *.jpeg *.png")
        ]
        files = filedialog.askopenfilenames(title="Seleccionar Archivos", filetypes=file_types)
        if files:
            self.selected_files = list(files)
            self.selected_files_label.configure(text=f"Archivos seleccionados: {len(files)}")
            self.update_file_listbox()
            self.move_up_button.configure(state=tk.NORMAL)
            self.move_down_button.configure(state=tk.NORMAL)
        else:
            self.selected_files_label.configure(text="No se han seleccionado archivos.")
            self.move_up_button.configure(state=tk.DISABLED)
            self.move_down_button.configure(state=tk.DISABLED)
        
        print("Archivos seleccionados:")
        for file in files:
            print(file)
    
    def update_file_listbox(self):
        """Actualiza la lista de archivos en la listbox."""

        self.file_listbox.delete(0, tk.END)
        for file in self.selected_files:
            self.file_listbox.insert(tk.END, os.path.basename(file))
    
    def move_up(self):
        """Mueve el archivo seleccionado hacia arriba en la lista."""

        selected = self.file_listbox.curselection()
        if selected is not None and selected > 0:
            index = selected
            self.selected_files[index-1], self.selected_files[index] = self.selected_files[index], self.selected_files[index-1]
            self.file_listbox.deselect(index)
            self.update_file_listbox()
            self.file_listbox.select(index-1)
            self.file_listbox._parent_canvas.yview_moveto((index-1) / self.file_listbox.size())  # Mantener la visualización del ítem seleccionado
    
    def move_down(self):
        """Mueve el archivo seleccionado hacia abajo en la lista."""

        selected = self.file_listbox.curselection()
        if selected is not None and selected < self.file_listbox.size() - 1:
            index = selected
            self.selected_files[index], self.selected_files[index+1] = self.selected_files[index+1], self.selected_files[index]
            self.file_listbox.deselect(index)

            self.update_file_listbox()
            self.file_listbox.select(index+1)
            self.file_listbox._parent_canvas.yview_moveto(index / self.file_listbox.size())  # Mantener la visualización del ítem seleccionado
    
    def volver_menu_principal(self):
        """restaura todos los valores y vuelve a menu principal"""
        self.master.show_frame("MainMenu")
        self.selected_files = []
        self.update_file_listbox()
        self.selected_files_label.configure(text="No se han seleccionado archivos.")
        self.move_up_button.configure(state=tk.DISABLED)
        self.move_down_button.configure(state=tk.DISABLED)
        # Resetear opciones de foliación
        self.foliar_var.set(False)
        self.start_page_var.set(1)
        self.start_page_entry.configure(state='disabled')
        
    def generar_expediente(self, files):
        # Process each file based on its type
        pdf_to_merge = []
        open_files = []  # Lista para mantener archivos abiertos

        for file in files:
            if file.lower().endswith('.docx'):
                word_to_pdf_file = self.convert_word_to_pdf(file)
                if word_to_pdf_file:
                    pdf_to_merge.append((word_to_pdf_file, file))
            elif file.lower().endswith(('.jpg', '.jpeg', '.png')):
                img_to_pdf_file = self.convert_image_to_pdf(file)
                if img_to_pdf_file:
                    pdf_to_merge.append((img_to_pdf_file, file))
            elif file.lower().endswith('.pdf'):
                f = open(file, 'rb')
                open_files.append(f)
                pdf_to_merge.append((f, file))
            else:
                print(f"Tipo de archivo no soportado: {file}")

        pdf_writer = PdfWriter()

        if self.foliar_var.get():
            try:
                start_page_number = int(self.start_page_var.get())
            except ValueError:
                messagebox.showerror("Error", "El número de página inicial debe ser un número entero.")
                return
        else:
            start_page_number = 1

        current_page = start_page_number - 1  # Adjust current page

        indice_data = []  # List to store document names and page ranges

        for pdf_file, original_file in pdf_to_merge:
            pdf_reader = PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)

            # Copiar páginas al escritor
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                if self.foliar_var.get():
                    page_number = current_page + page_num + 1
                    page = self.add_page_number(page, page_number)
                pdf_writer.add_page(page)

            # Ajustar y copiar marcadores
            if pdf_reader.outline:
                self._add_outline(pdf_writer, pdf_reader, pdf_reader.outline, current_page)

            # Keep track of page ranges for the Indice
            start_page = current_page + 1
            end_page = current_page + num_pages
            document_name = os.path.basename(original_file)
            indice_data.append((document_name, start_page, end_page))

            current_page += num_pages

        if files:
            output_dir = os.path.dirname(files[0])
        else:
            output_dir = os.getcwd()

        output_filename = os.path.join(output_dir, 'expediente.pdf')
        counter = 1
        while os.path.exists(output_filename):
            output_filename = os.path.join(output_dir, f'expediente_{counter}.pdf')
            counter += 1

        with open(output_filename, 'wb') as out_file:
            pdf_writer.write(out_file)

        # Generate Indice Word document if foliation is activated
        if self.foliar_var.get():
            self.generar_indice(indice_data)

        # Mostrar mensaje de éxito
        messagebox.showinfo("Expediente Creado", "El expediente ha sido creado con éxito.")

        # Cerrar archivos abiertos
        for f in open_files:
            f.close()

        # Eliminar archivos convertidos
        for file in files:
            if file.lower().endswith(('.docx', '.jpg', '.jpeg', '.png')):
                converted_file = file.rsplit('.', 1)[0] + '_converted.pdf'
                if os.path.exists(converted_file):
                    os.remove(converted_file)
                    print(f"Archivo eliminado: {converted_file}")

    def _add_outline(self, pdf_writer, pdf_reader, outlines, current_page, parent=None):
        for outline in outlines:
            if isinstance(outline, list):
                self._add_outline(pdf_writer, pdf_reader, outline, current_page, parent)
            elif isinstance(outline, PyPDF2.generic.Destination):
                page_number_in_pdf = pdf_reader.get_destination_page_number(outline)
                pdf_writer.add_outline_item(
                    outline.title,
                    page_number_in_pdf + current_page,
                    parent=parent
                )
            elif isinstance(outline, PyPDF2.generic.OutlineItem):
                if outline.destination:
                    page_number_in_pdf = pdf_reader.get_destination_page_number(outline.destination)
                    new_parent = pdf_writer.add_outline_item(
                        outline.title,
                        page_number_in_pdf + current_page,
                        parent=parent
                    )
                else:
                    new_parent = pdf_writer.add_outline_item(
                        outline.title,
                        current_page,
                        parent=parent
                    )
                if outline.children:
                    self._add_outline(pdf_writer, pdf_reader, outline.children, current_page, new_parent)

    def _recorrer_outlines(self, outlines, all_outlines, current_page):
        for outline in outlines:
            if isinstance(outline, list):
                self._recorrer_outlines(outline, all_outlines, current_page)
            elif isinstance(outline, PyPDF2.generic.Destination):
                all_outlines.append({
                    'title': outline.title,
                    'page': outline.page_number + current_page
                })
            elif isinstance(outline, PyPDF2.generic.OutlineItem):
                self._recorrer_outlines([outline], all_outlines, current_page)

    def add_page_number(self, page, number):
        # Obtener el tamaño de la página
        page_width = float(page.mediabox.upper_right[0]) - float(page.mediabox.lower_left[0])
        page_height = float(page.mediabox.upper_right[1]) - float(page.mediabox.lower_left[1])

        packet = BytesIO()
        can = canvas.Canvas(packet, pagesize=(page_width, page_height))

        # Ajustar las coordenadas para la esquina superior derecha
        x_position = page_width - 50  # Ajusta este valor si es necesario
        y_position = page_height - 50  # Ajusta este valor si es necesario

        can.drawString(x_position, y_position, str(number))
        can.save()

        packet.seek(0)
        new_pdf = PdfReader(packet)
        page.merge_page(new_pdf.pages[0])
        return page

    def convert_word_to_pdf(self, file_path):
        output_path = file_path.rsplit('.', 1)[0] + '_converted' + '.pdf'
        try:
            convert(file_path, output_path)
            print(f"Convirtiendo Word a PDF: {file_path}")
        except Exception as e:
            print(f"Sin Microsoft Word instalado. Ejecutando con LibreOffice")
            from subprocess import Popen
            LIBRE_OFFICE = r"C:\Program Files\LibreOffice\program\soffice.exe"
            out_folder = os.path.dirname(file_path)
            p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
                        out_folder, file_path])
            p.communicate()
            libre_office_output = file_path.rsplit('.', 1)[0] + '.pdf'
            if os.path.exists(libre_office_output):
                os.rename(libre_office_output, output_path)
                print(f"Archivo convertido con LibreOffice: {output_path}")
            else:
                print(f"No se encontró el archivo convertido por LibreOffice: {libre_office_output}")
        return output_path
        
    def convert_image_to_pdf(self, file_path):
        try:
            image = Image.open(file_path)
            output_path = file_path.rsplit( '.', 1)[0] + '_converted' + '.pdf'
            image.save(output_path, 'PDF', resolution=100.0)
            print(f"Convirtiendo Imagen a PDF: {file_path}")
            return output_path
        except Exception as e:
            print(f"Error al convertir {file_path}: {e}")
            return None

    def generar_indice(self, indice_data):
        # Create a new Word document
        
        document = Document('templates/indice_template.docx')

        # Assuming the first table is the target table
        table = document.tables[0]

        for doc_name, start_page, end_page in indice_data:
            row_cells = table.add_row().cells
            row_cells[0].text = doc_name
            for paragraph in row_cells[0].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if start_page == end_page:
                row_cells[1].text = f"{start_page}"
            else:
                row_cells[1].text = f"{start_page}-{end_page}"
            for paragraph in row_cells[1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save the Indice document in the same directory as the original files
        if self.selected_files:
            output_dir = os.path.dirname(self.selected_files[0])
        else:
            output_dir = os.getcwd()

        indice_filename = os.path.join(output_dir, 'Indice.docx')
        counter = 1
        while os.path.exists(indice_filename):
            indice_filename = os.path.join(output_dir, f'Indice_{counter}.docx')
            counter += 1

        document.save(indice_filename)
        print(f"Indice generado: {indice_filename}")

    def toggle_start_page_entry(self):
        if self.foliar_var.get():
            self.start_page_entry.configure(state='normal')
        else:
            self.start_page_entry.configure(state='disabled')


